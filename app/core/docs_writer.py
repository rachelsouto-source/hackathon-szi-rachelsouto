"""
Geração das saídas da DD:
- Parecer em Markdown (a partir do resultado da auditoria + template).
- Google Doc na pasta do empreendimento (upload de HTML convertido p/ Google Doc — só Drive scope).
- Planilha .xlsx de controle (openpyxl), retornada em bytes.
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any

# ---------- Parecer (Markdown) ----------

def _tabela_areas(titulo: str, linhas: list[dict], col_ref: str = "Referência") -> str:
    if not linhas:
        return ""
    out = [f"**{titulo}**", "", f"| {col_ref} | Área (m²) |", "|---|---|"]
    for ln in linhas:
        out.append(f"| {ln.get('ref','')} | {ln.get('area','')} |")
    return "\n".join(out) + "\n"


_GRUPOS = [
    ("IMÓVEL", ["matrícula", "matricula", "cadastral", "confrontante", "viabilidade"]),
    ("ESTUDOS", ["topográf", "topograf", "ambiental", "eva", "sondagem", "estrutura", "fundaç", "fundac"]),
    ("VALIDAÇÃO", ["validação", "validacao"]),
    ("NEGÓCIO / LOCALIZAÇÃO", ["proposta", "imagens", "localiza", "spu", "instalaç", "instalac"]),
]


def _docs_por_grupo(achados: list[dict]) -> str:
    blocos = []
    for titulo, chaves in _GRUPOS:
        itens = []
        for a in achados:
            et = (a.get("etapa", "") or "").lower()
            if any(k in et for k in chaves):
                links = [l for l in [a.get("link"), a.get("link2")] if l]
                urls = " · ".join(f"[abrir]({l})" for l in links) if links else "*(sem link)*"
                itens.append(f"- {a.get('etapa','')}: {urls}")
        if itens:
            blocos.append(f"**{titulo}**\n" + "\n".join(itens))
    return "\n\n".join(blocos) or "- (fontes serão preenchidas a partir do Drive)"


def render_parecer_md(nome: str, r: dict[str, Any]) -> str:
    im = r.get("imovel", {})
    prop = r.get("proprietarios", [])
    con = r.get("conclusao", {})
    neg = r.get("negocio", {})
    val = r.get("validacao", {})
    at = r.get("areas_tabela", {})
    flags = "\n".join(f"  {i+1}. {f}" for i, f in enumerate(neg.get("red_flags", []))) or "  —"

    tabelas = "\n".join(filter(None, [
        _tabela_areas("Área de Matrícula", at.get("matricula", []), "Referência (Inscrição Imobiliária)"),
        _tabela_areas("Área de Cadastro Imobiliário PMF", at.get("cadastro_pmf", []), "Referência (Inscrição Imobiliária)"),
        (f"**Área Levantamento Topográfico**\n\n| Referência | Área (m²) |\n|---|---|\n| Área Real | {at.get('topografico','')} |\n"
         if at.get("topografico") else ""),
    ]))

    def _lista(itens, vazio="—"):
        return "\n".join(f"- {x}" for x in itens) if itens else f"- {vazio}"

    ajustes = ("\n\n" + _lista(val.get("ajustes", []))) if val.get("ajustes") else ""
    docs_ap = ("\n\n**Documentos para Aprovação do Projeto Arquitetônico (PMF):**\n"
               + _lista(val.get("docs_aprovacao", []))) if val.get("docs_aprovacao") else ""
    docs_al = ("\n\n**Documentos para o Alvará de Construção:**\n"
               + _lista(val.get("docs_alvara", []))) if val.get("docs_alvara") else ""

    # LOCALIZAÇÃO: imagens reais (drone/localização) — entram ANTES de tudo
    figs = r.get("figuras", [])
    localizacao_md = "### LOCALIZAÇÃO\nVisão geral de onde o empreendimento (modelo Spot) está inserido:\n\n"
    if figs:
        localizacao_md += "\n".join(f"![{f.get('cap','')}]({f.get('url','')})" for f in figs) + "\n"
    else:
        localizacao_md += "::FIG:: Localização / entorno do terreno\n"

    # figura por seção = QUADRO/placeholder formatado (a usuária insere a imagem do tópico)
    def figof(key_cap: str) -> str:
        return f"\n\n::FIG:: {key_cap}"

    return f"""# PARECER TÉCNICO – DUE DILIGENCE – {nome.upper()}

> Gerado pela aplicação Auditor de DD Técnica. Requer revisão humana.

## 1. IMÓVEL

| Item | Descrição |
|---|---|
| **Inscrição** | {im.get('inscricoes','(pendente)')} |
| **Endereço** | {im.get('endereco','(pendente)')} |
| **Área total** | {im.get('area_matricula_total','(pendente)')} |
| **Matrícula** | {im.get('matriculas','(pendente)')} |

## 2. PROPRIETÁRIO(A)
{chr(10).join('- ' + p for p in prop) if prop else '- (pendente)'}

Trata-se de parecer técnico acerca das diligências e da análise técnica realizada pelo Setor de Lançamentos da Seazone Investimentos em relação aos estudos de viabilidade referentes ao terreno e ao estudo preliminar desenvolvido, objetivando a aquisição do imóvel e continuidade no processo de estruturação do empreendimento.

### Documentos analisados (fontes)
{_docs_por_grupo(r.get('achados', []))}

Para realização da due diligence, foi verificado, entre outros tópicos, código de obras, plano diretor e demais legislações vigentes.

## 3. CONCLUSÃO

{localizacao_md}
### TOPOGRAFIA
{con.get('topografia','(pendente)')}

{tabelas}{figof('Levantamento topográfico')}

### ESTUDO PRÉVIO AMBIENTAL
{con.get('ambiental','(pendente)')}{figof('Mapa de restrições ambientais / EVA')}

### VALIDAÇÃO DO ESTUDO PRELIMINAR SEAZONE
{con.get('validacao_ep','(pendente)')}{ajustes}{docs_ap}{docs_al}{figof('Estudo preliminar / ajustes do anteprojeto')}

### SONDAGEM
{con.get('sondagem','(pendente)')}{figof('Locação dos furos de sondagem')}

### ESTRUTURA / FUNDAÇÃO
{con.get('estrutura_fundacao','(pendente)')}{figof('Implantação estrutural / locação das estacas')}

### LEITURA DE NEGÓCIO *(complemento do auditor)*
- **Impacto em custo/prazo:** {neg.get('impacto_custo_prazo','—')}
- **Red flags priorizados:**
{flags}
- **Aproveitamento × VGV (estimado):** {neg.get('aproveitamento_vgv','—')}
- **Recomendação:** {neg.get('recomendacao','—')}

### CONCLUSÃO
{con.get('final','(pendente)')}

*Florianópolis/SC, {date.today().strftime('%d/%m/%Y')}.*
*Setor de Projetos — Estruturação — Seazone Investimentos.*
"""


def _inline(s: str) -> str:
    import re
    # imagens ![alt](url)  (antes dos links)
    s = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)",
               r'<img src="\2" alt="\1" style="max-width:480px"/><br/><i>\1</i>', s)
    # links [texto](url)
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', s)
    # negrito **x**
    while "**" in s:
        s = s.replace("**", "<b>", 1).replace("**", "</b>", 1)
    return s


def _md_to_html(md: str) -> str:
    import re
    html = ["<html><body style=\"font-family:Arial,sans-serif;font-size:11pt\">"]
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].rstrip()
        if s.startswith("|"):  # tabela
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(lines[i].strip()); i += 1
            html.append('<table border="1" cellspacing="0" cellpadding="5" style="border-collapse:collapse">')
            for ri, row in enumerate(rows):
                if re.match(r"^\|[\s:|-]+\|?$", row):
                    continue
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                tag = "th" if ri == 0 else "td"
                html.append("<tr>" + "".join(f"<{tag}>{_inline(c)}</{tag}>" for c in cells) + "</tr>")
            html.append("</table>")
            continue
        if s.startswith("### "):
            txt = s[4:]
            html.append(f'<h3 style="color:#2f5597;border-bottom:1px solid #d0d7e6;'
                        f'padding-bottom:3px">{_inline(txt)}</h3>')
        elif s.startswith("## "):
            # faixa colorida numerada (estilo template SZI)
            txt = s[3:]
            cor = "#ff6b5e" if "CONCLUS" in txt.upper() else "#0f1e3d"
            html.append(
                f'<table width="100%" style="border-collapse:collapse;margin:14px 0 6px">'
                f'<tr><td style="background-color:{cor};color:#ffffff;padding:7px 12px;'
                f'font-weight:bold;font-size:13pt;letter-spacing:.5px">{_inline(txt)}</td></tr></table>')
        elif s.startswith("# "):
            html.append(f'<h1 style="text-align:center;color:#0f1e3d">{_inline(s[2:])}</h1>')
        elif s.startswith("> "):
            html.append(f"<blockquote>{_inline(s[2:])}</blockquote>")
        elif s.startswith("- "):
            html.append(f"<li>{_inline(s[2:])}</li>")
        elif s == "---":
            html.append("<hr/>")
        elif s:
            html.append(f"<p>{_inline(s)}</p>")
        i += 1
    html.append("</body></html>")
    return "\n".join(html)


def _docx_shade(paragraph, fill_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _docx_hyperlink(paragraph, url: str, text: str, color="2F5597"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _docx_runs(paragraph, text: str) -> None:
    """Adiciona runs interpretando **negrito**, [texto](url) e ![cap](url) (como link 'ver figura')."""
    import re
    # imagem -> linha de link (sem imagem quebrada)
    m_img = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", text.strip())
    if m_img:
        paragraph.add_run("📎 Figura — " + m_img.group(1) + "  ").italic = True
        _docx_hyperlink(paragraph, m_img.group(2), "ver figura")
        return
    token = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for mt in token.finditer(text):
        if mt.start() > pos:
            paragraph.add_run(text[pos:mt.start()])
        tk = mt.group(0)
        if tk.startswith("**"):
            paragraph.add_run(tk[2:-2]).bold = True
        else:
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tk)
            _docx_hyperlink(paragraph, mm.group(2), mm.group(1))
        pos = mt.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def gerar_docx_bytes(nome: str, parecer_md: str, images: dict | None = None) -> bytes:
    """Gera o parecer como .docx formatado. `images` = {url: bytes} p/ embutir imagens reais
    (ex.: localização); demais figuras viram quadro/placeholder formatado."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    images = images or {}

    doc = Document()
    # Fonte padrão: Calibri 12 (igual ao template da SZI)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(12)
    # garante que o tema de fonte também use Calibri
    rFonts = style_normal.element.rPr.get_or_add_rFonts() if hasattr(style_normal.element, 'rPr') else None

    # ---- Logo Seazone no topo ----
    import pathlib as _pl
    _logo_path = _pl.Path(__file__).resolve().parent.parent / "static" / "seazone-logo.png"
    if _logo_path.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_p.add_run()
        logo_run.add_picture(str(_logo_path), width=Inches(2.0))
        doc.add_paragraph()  # espaço após logo

    def _placeholder(cap: str):
        tb = doc.add_table(rows=1, cols=1)
        tb.style = "Table Grid"
        cp = tb.rows[0].cells[0].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_shade(cp, "F2F4F8")
        rr = cp.add_run("📷  Inserir imagem aqui"); rr.italic = True
        rr.font.color.rgb = RGBColor(0x8A, 0x93, 0xA6)
        tb.rows[0].cells[0].add_paragraph(); tb.rows[0].cells[0].add_paragraph()
        legp = doc.add_paragraph(); legp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = legp.add_run(f"Figura — {cap}"); lr.italic = True; lr.bold = True
        lr.font.size = Pt(10); lr.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    def _image(cap: str, data: bytes):
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(data), width=Inches(5.5))
            legp = doc.add_paragraph(); legp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = legp.add_run(f"Figura — {cap}"); lr.italic = True; lr.bold = True
            lr.font.size = Pt(10); lr.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)
        except Exception:  # noqa: BLE001
            _placeholder(cap)

    lines = parecer_md.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].rstrip()
        if s.startswith("|"):  # tabela
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(lines[i].strip()); i += 1
            data = []
            for ri, row in enumerate(rows):
                import re as _re
                if _re.match(r"^\|[\s:|-]+\|?$", row):
                    continue
                data.append([c.strip() for c in row.strip().strip("|").split("|")])
            if data:
                t = doc.add_table(rows=0, cols=len(data[0]))
                t.style = "Table Grid"
                for ri, cells in enumerate(data):
                    crow = t.add_row().cells
                    for ci, cval in enumerate(cells):
                        if ci < len(crow):
                            p = crow[ci].paragraphs[0]
                            _docx_runs(p, cval)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
            continue
        if s.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s[2:]); run.bold = True; run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3D)
        elif s.startswith("## "):
            import re as _re
            mh = _re.match(r"(\d+)\.\s*(.+)", s[3:])
            num = mh.group(1) if mh else ""
            title = (mh.group(2) if mh else s[3:]).upper()
            is_concl = "CONCLUS" in title
            WHITE = RGBColor(0xFF, 0xFF, 0xFF)
            tb = doc.add_table(rows=1, cols=2 if num else 1)
            tb.autofit = False
            cells = tb.rows[0].cells
            if num:
                cells[0].width = Inches(0.45)
                cp0 = cells[0].paragraphs[0]; cp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _docx_shade(cp0, "0B1530")
                r0 = cp0.add_run(num); r0.bold = True; r0.font.color.rgb = WHITE; r0.font.size = Pt(13)
                tcell = cells[1]; tcell.width = Inches(6.05)
            else:
                tcell = cells[0]
            cpt = tcell.paragraphs[0]
            _docx_shade(cpt, "FF6B5E" if is_concl else "16284D")
            rt = cpt.add_run("  " + title); rt.bold = True; rt.font.color.rgb = WHITE; rt.font.size = Pt(13)
        elif s.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(s[4:].upper()); run.bold = True; run.font.size = Pt(12); run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3D)
        elif s.startswith("> "):
            p = doc.add_paragraph(); r = p.add_run(s[2:]); r.italic = True
            r.font.color.rgb = RGBColor(0x69, 0x74, 0x8C)
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); _docx_runs(p, s[2:])
        elif s.startswith("::FIG:: "):
            _placeholder(s[8:].strip())
        elif s.startswith("!["):
            import re as _re
            m = _re.match(r"!\[([^\]]*)\]\(([^)]*)\)", s)
            cap = (m.group(1) if m else "").replace("Figura — ", "").strip()
            url = m.group(2) if m else ""
            if url in images:
                _image(cap, images[url])
            else:
                _placeholder(cap)
        elif s == "---":
            doc.add_paragraph()
        elif s.strip():
            p = doc.add_paragraph(); _docx_runs(p, s)
        i += 1

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def create_google_doc(folder_id: str, title: str, parecer_md: str, images: dict | None = None) -> dict:
    """Cria um Google Doc na pasta convertendo um .docx formatado. Requer Service Account."""
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload
    from .drive_client import _credentials

    svc = build("drive", "v3", credentials=_credentials(), cache_discovery=False)
    docx_bytes = gerar_docx_bytes(title, parecer_md, images)
    media = MediaIoBaseUpload(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False)
    meta = {
        "name": title,
        "mimeType": "application/vnd.google-apps.document",  # converte o .docx para Google Doc
        "parents": [folder_id],
    }
    f = svc.files().create(
        body=meta, media_body=media,
        fields="id, webViewLink", supportsAllDrives=True,
    ).execute()
    return {"id": f["id"], "url": f.get("webViewLink", "")}


# ---------- Planilha de controle (.xlsx) ----------

_FILL = {"OK": "C6EFCE", "Pendente": "FFEB9C", "Divergência": "FFC7CE",
         "Divergencia": "FFC7CE", "Não se aplica": "D9D9D9", "Nao se aplica": "D9D9D9"}
_FONT = {"OK": "006100", "Pendente": "9C6500", "Divergência": "9C0006",
         "Divergencia": "9C0006", "Não se aplica": "595959", "Nao se aplica": "595959"}
_HEADERS = ["#", "Etapa", "Documento", "Status", "Severidade", "Observação", "Ação recomendada", "Fonte"]


_OBRIGATORIOS = ["Matrícula", "Certidão cadastral", "Certidão de confrontantes",
                 "Viabilidade", "Levantamento topográfico", "Estudo ambiental",
                 "Sondagem", "Estrutura", "Fundação", "Validação do EP"]


def gerar_xlsx_bytes(nome: str, achados: list[dict[str, Any]],
                     recomendacao: str = "—", data_str: str = "") -> bytes:
    from collections import Counter
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    data_str = data_str or date.today().strftime("%d/%m/%Y")
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    link_font = Font(color="2F5597", underline="single")

    wb = Workbook()

    # ---------- Aba 1: Controle ----------
    ws = wb.active
    ws.title = "Controle DD Técnica"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:H1")
    c = ws["A1"]
    c.value = f"Controle DD Técnica — {nome}"
    c.font = Font(bold=True, size=14, color="FFFFFF")
    c.fill = PatternFill("solid", fgColor="0F1E3D")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # linha de info (empreendimento / data / recomendação)
    reco_fill = ("C6EFCE" if recomendacao.upper().startswith("GO") and "NO-GO" not in recomendacao.upper()
                 else "FFC7CE" if "NO-GO" in recomendacao.upper() else "FFEB9C")
    ws.merge_cells("A2:C2"); ws["A2"] = f"Empreendimento: {nome}"
    ws.merge_cells("D2:E2"); ws["D2"] = f"Data: {data_str}"
    ws.merge_cells("F2:H2"); ws["F2"] = f"Recomendação: {recomendacao}"
    ws["A2"].font = ws["D2"].font = Font(bold=True)
    ws["F2"].font = Font(bold=True, color="000000")
    ws["F2"].fill = PatternFill("solid", fgColor=reco_fill)
    for cc in ("A2", "D2", "F2"):
        ws[cc].alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[2].height = 20

    hdr = 3
    for col, h in enumerate(_HEADERS, start=1):
        cell = ws.cell(row=hdr, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F5597")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    for i, item in enumerate(achados, start=1):
        rr = hdr + i
        status = item.get("status", "Pendente")
        vals = [i, item.get("etapa", ""), item.get("documento", ""), status,
                item.get("severidade", ""), item.get("observacao", ""),
                item.get("acao", ""), item.get("fonte", "")]
        for col, val in enumerate(vals, start=1):
            cell = ws.cell(row=rr, column=col, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
        # status colorido
        sc = ws.cell(row=rr, column=4)
        if status in _FILL:
            sc.fill = PatternFill("solid", fgColor=_FILL[status])
            sc.font = Font(bold=True, color=_FONT.get(status, "000000"))
            sc.alignment = Alignment(horizontal="center", vertical="center")
        # fonte com hyperlink (se houver link)
        link = item.get("link") or ""
        if link:
            fc = ws.cell(row=rr, column=8)
            fc.hyperlink = link
            fc.font = link_font

    nrows = hdr + len(achados)
    # filtro automático + congelar
    ws.auto_filter.ref = f"A{hdr}:H{nrows}"
    ws.freeze_panes = "A4"
    # lista suspensa no Status (para atualização manual)
    dv = DataValidation(type="list",
                        formula1='"OK,Pendente,Divergência,Não se aplica"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"D{hdr+1}:D{nrows}")
    for col, w in enumerate([4, 24, 26, 14, 12, 50, 42, 30], start=1):
        ws.column_dimensions[get_column_letter(col)].width = w

    # ---------- Aba 2: Resumo ----------
    ws2 = wb.create_sheet("Resumo")
    ws2.sheet_view.showGridLines = False
    ws2["A1"] = f"Resumo — {nome}"
    ws2["A1"].font = Font(bold=True, size=13, color="FFFFFF")
    ws2["A1"].fill = PatternFill("solid", fgColor="0F1E3D")
    ws2.merge_cells("A1:B1")
    ws2["A3"] = "Recomendação"; ws2["B3"] = recomendacao
    ws2["A3"].font = Font(bold=True); ws2["B3"].font = Font(bold=True)
    ws2["B3"].fill = PatternFill("solid", fgColor=reco_fill)
    cnt = Counter(a.get("status", "Pendente") for a in achados)
    row = 5
    ws2.cell(row=row, column=1, value="Status").font = Font(bold=True)
    ws2.cell(row=row, column=2, value="Qtd").font = Font(bold=True)
    for st in ["OK", "Pendente", "Divergência", "Não se aplica"]:
        row += 1
        ws2.cell(row=row, column=1, value=st)
        ws2.cell(row=row, column=2, value=cnt.get(st, 0))
        if st in _FILL:
            ws2.cell(row=row, column=1).fill = PatternFill("solid", fgColor=_FILL[st])
    row += 1
    ws2.cell(row=row, column=1, value="Total").font = Font(bold=True)
    ws2.cell(row=row, column=2, value=len(achados)).font = Font(bold=True)

    # checklist de documentos obrigatórios
    row += 2
    ws2.cell(row=row, column=1, value="Documento obrigatório").font = Font(bold=True)
    ws2.cell(row=row, column=2, value="Presente?").font = Font(bold=True)
    etapas_txt = " | ".join((a.get("etapa", "") or "") for a in achados).lower()
    for doc in _OBRIGATORIOS:
        row += 1
        presente = doc.lower() in etapas_txt
        ws2.cell(row=row, column=1, value=doc)
        cc = ws2.cell(row=row, column=2, value="Sim" if presente else "FALTA")
        cc.fill = PatternFill("solid", fgColor="C6EFCE" if presente else "FFC7CE")
        cc.font = Font(bold=True, color="006100" if presente else "9C0006")
    ws2.column_dimensions["A"].width = 34
    ws2.column_dimensions["B"].width = 14

    # ---------- Aba 3: Ações / Pendências ----------
    ws3 = wb.create_sheet("Ações e Pendências")
    ws3.sheet_view.showGridLines = False
    ws3["A1"] = f"Ações e Pendências — {nome}"
    ws3["A1"].font = Font(bold=True, size=13, color="FFFFFF")
    ws3["A1"].fill = PatternFill("solid", fgColor="0F1E3D")
    ws3.merge_cells("A1:D1")
    heads3 = ["Etapa", "Status", "Ação recomendada", "Fonte"]
    for col, h in enumerate(heads3, start=1):
        cell = ws3.cell(row=2, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F5597")
        cell.border = border
    r3 = 2
    for a in achados:
        if a.get("status") in ("Pendente", "Divergência", "Divergencia"):
            r3 += 1
            vals = [a.get("etapa", ""), a.get("status", ""), a.get("acao", ""), a.get("fonte", "")]
            for col, val in enumerate(vals, start=1):
                cell = ws3.cell(row=r3, column=col, value=val)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border
            sc = ws3.cell(row=r3, column=2)
            stt = a.get("status", "")
            if stt in _FILL:
                sc.fill = PatternFill("solid", fgColor=_FILL[stt])
                sc.font = Font(bold=True, color=_FONT.get(stt, "000000"))
            if a.get("link"):
                fc = ws3.cell(row=r3, column=4); fc.hyperlink = a["link"]; fc.font = link_font
    if r3 == 2:
        ws3.cell(row=3, column=1, value="Sem pendências 🎉")
    for col, w in enumerate([24, 14, 50, 30], start=1):
        ws3.column_dimensions[get_column_letter(col)].width = w
    ws3.freeze_panes = "A3"

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()
